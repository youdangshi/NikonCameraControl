# -*- coding: utf-8 -*-
"""Nikon Camera Control 商业项目介绍文档生成脚本（standard_business_brief 预设）。"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\李\Desktop\项目组\nikon-camera-control\docs\Nikon_Camera_Control_项目介绍.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)


def set_run(run, size=11, color="1A1A1A", bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
):
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.10

title = styles["Title"]
title.font.name = "Calibri"
title.font.size = Pt(24)
title.font.bold = True
title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(6)


def heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for r in p.runs:
        set_run(r, size={1: 16, 2: 13, 3: 12}[level], color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}[level], bold=True)
    return p


def body(text, size=11, bold=False, color="1A1A1A"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r, size=10.5)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=10.5)
    return p


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)
    total = sum(int(round(w * 1440)) for w in widths)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tblPr.append(ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)
    grid = tbl.find(qn("w:tblGrid"))
    for gc in list(grid):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for el in tcPr.findall(qn("w:tcW")):
                tcPr.remove(el)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(widths[i] * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def table(headers, rows, widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, size=10, bold=True, color="FFFFFF")
        shade(cell, "2E74B5")
        p.paragraph_format.space_after = Pt(0)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run(r, size=10)
            p.paragraph_format.space_after = Pt(0)
            if i % 2 == 0:
                shade(cell, "F2F4F7")
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def footer_page():
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run("第 ")
    set_run(r1, size=9, color="585870")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    rnd = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    rnd.append(t)
    fld.append(rnd)
    p._p.append(fld)
    r2 = p.add_run(" 页")
    set_run(r2, size=9, color="585870")


footer_page()

# ── 标题区 ────────────────────────────────────────────
p = doc.add_paragraph(style="Title")
r = p.add_run("Nikon Camera Control")
set_run(r, size=24, color="0B2545", bold=True)
p2 = doc.add_paragraph()
r2 = p2.add_run("产品与商业项目介绍 · 商业展示版")
set_run(r2, size=11, color="585870")
p3 = doc.add_paragraph()
r3 = p3.add_run("面向合作伙伴、新团队成员与潜在投资者的项目说明。")
set_run(r3, size=11, bold=True, color="1F3A5F")

heading("1. 文档目的与位置")
table(
    ["文档", "位置", "用途"],
    [
        ["本介绍文档", "docs/Nikon_Camera_Control_项目介绍.docx", "对外商业展示、团队认知、投资/合作说明"],
        ["技术交接文档", "docs/Nikon_Camera_Control_交接文档.docx", "新对话/新开发者从零快速接收项目"],
        ["开源上传清单", "docs/OPEN_SOURCE.md", "明确 GitHub 可上传与禁止上传内容"],
        ["测试 APK", "app/release/NikonCameraControl-1.0.0-debug.apk", "真机安装与验证"],
    ],
    [1.35, 2.65, 2.5],
)

heading("2. 产品定位")
body("Nikon Camera Control 是一款面向 Nikon Z 系列用户的独立 Android 相机控制与照片同步工具，通过 STA、WiFi 热点和 USB 三种方式连接相机，并在手机上完成遥控拍摄、参数调节、照片浏览、AI 修图建议与人像构图引导。")
bullet("核心价值：让摄影者用一块手机屏幕完成从取景、摆姿、拍摄到修图的完整工作流。")
bullet("差异化：不依赖电脑中转，手机可直连相机；AI 模型可切换；内置可调透明度的姿势框线。")

heading("3. 目标用户与场景")
bullet("Vlog 创作者：需要手机边看边拍、快速调整参数。")
bullet("人像摄影师：需要构图引导与轻量 AI 后期建议。")
bullet("相机用户：需要无线/有线快速同步照片到手机。")
bullet("工作室/团队：希望低成本搭建可复用的相机辅助工具。")

heading("4. 核心功能")
table(
    ["功能", "说明"],
    [
        ["PTP/IP 无线遥控", "主路径，手机/电脑直连相机 192.168.1.1:15740"],
        ["STA 局域网", "相机加入家庭 WiFi，同网输入相机 IP 后连接"],
        ["USB 有线备选", "电脑可走 WinUSB；手机端后续接入 OTG 通道"],
        ["参数控制", "ISO、快门、光圈、曝光补偿、白平衡、对焦"],
        ["实时取景", "Live View 与自动对焦"],
        ["照片同步", "相机照片列表、本地媒体、同步进度"],
        ["AI 修图", "支持 DeepSeek / OpenAI 兼容模型，可传图分析"],
        ["人像姿势框线", "半身、全身、头部、三分、网格，透明度可调"],
    ],
    [1.7, 4.8],
)

heading("5. 当前开发进度")
table(
    ["模块", "进度", "说明"],
    [
        ["Android APK 构建", "可打包", "已有可安装 debug APK"],
        ["手机原生 TCP", "已完成", "直连相机 PTP/IP 端口"],
        ["UI 五个主页面", "已完成", "首页/我的相机/相机照片/同步/本地照片"],
        ["三种连接入口", "界面完成", "STA/WiFi 可用；USB 手机端待原生接入"],
        ["AI 修图", "已完成", "模型/地址/名称可配置"],
        ["姿势框线", "已完成", "已接入实时取景与设置"],
        ["真机协议验证", "待办", "需用相机日志校准 PTP/IP 握手"],
    ],
    [1.55, 1.35, 3.6],
)

heading("6. 商业价值与变现方向")
bullet("免费版：基础连接、遥控、姿势框线、基础同步。")
bullet("专业版：AI 高级修图、批量同步、多相机、无广告。")
bullet("团队/商用授权：定制机型适配、私有部署、接口接入。")
bullet("开源策略：核心源码开源，商业增强功能保留授权，建立社区与贡献者。")

heading("7. 短期路线图")
numbered("用真机验证 PTP/IP 握手并校准协议。")
numbered("接入手机端 Live View 数据阶段。")
numbered("实现 Android USB OTG 原生通道。")
numbered("完善真实照片下载、本地相册与同步队列。")
numbered("发布签署 Release APK 并启动 GitHub 开源。")

heading("8. 技术底座")
table(
    ["层", "技术", "说明"],
    [
        ["UI/App", "React + Capacitor", "跨平台移动端与 WebView 容器"],
        ["连接", "Node net / Capacitor TCP / libusb", "桌面、手机、USB 三态传输"],
        ["协议", "PTP/IP 纯 JS/TypeScript 实现", "命令、响应、事件与厂商扩展"],
        ["AI", "OpenAI 兼容接口", "可切换 DeepSeek / 自定义模型"],
        ["构建", "Gradle + electron-builder", "Android APK 与桌面 EXE"],
    ],
    [1.25, 2.6, 2.65],
)

heading("9. 交接说明")
body("若需要继续开发，请使用 docs/Nikon_Camera_Control_交接文档.docx 中的交接提示词，并遵循 docs/OPEN_SOURCE.md 的上传规则。当前版本是可安装测试版，正式发布前需要完成真机协议验证与 Release 签名。")

doc.save(OUT)
print("saved:", OUT)
