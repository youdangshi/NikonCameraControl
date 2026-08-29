# -*- coding: utf-8 -*-
"""Nikon Camera Control 项目交接文档生成脚本（compact_reference_guide 预设）。"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\李\Desktop\项目组\nikon-camera-control\docs\Nikon_Camera_Control_交接文档.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)


def set_run(run, size=11, color="1A1A1A", bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
):
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.25

title = styles["Title"]
title.font.name = "Calibri"
title.font.size = Pt(24)
title.font.bold = True
title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(6)


def heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for r in p.runs:
        set_run(r, size={1: 16, 2: 13, 3: 12}[level], color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}[level], bold=True)
    return p


def body(text, size=11, bold=False, color="1A1A1A"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.187)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run(r, size=10.5)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.187)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=10.5)
    return p


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)
    total = sum(int(round(w * 1440)) for w in widths)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tblPr.append(ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)
    grid = tbl.find(qn("w:tblGrid"))
    for gc in list(grid):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for el in tcPr.findall(qn("w:tcW")):
                tcPr.remove(el)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(widths[i] * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def table(headers, rows, widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, size=10, bold=True, color="FFFFFF")
        shade(cell, "2E74B5")
        p.paragraph_format.space_after = Pt(0)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run(r, size=10)
            p.paragraph_format.space_after = Pt(0)
            if i % 2 == 0:
                shade(cell, "F2F4F7")
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def footer_page():
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    set_run(r, size=9, color="585870")
    r1 = p.add_run("第 ")
    set_run(r1, size=9, color="585870")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    rnd = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    rnd.append(t)
    fld.append(rnd)
    p._p.append(fld)
    r2 = p.add_run(" 页")
    set_run(r2, size=9, color="585870")


footer_page()

# ── 标题区 ────────────────────────────────────────────
p = doc.add_paragraph(style="Title")
r = p.add_run("Nikon Camera Control 项目交接文档")
set_run(r, size=24, color="0B2545", bold=True)
p2 = doc.add_paragraph()
r2 = p2.add_run("版本 1.0  ·  交接日期 2026-08-28  ·  用途：交接给新对话 / 新开发者")
set_run(r2, size=11, color="585870")
p3 = doc.add_paragraph()
r3 = p3.add_run("本文件是可执行的交接说明书：先读它，再动手改代码。")
set_run(r3, size=11, bold=True, color="1F3A5F")

heading("0. 从零快速接收清单")
numbered("先读 nikon-camera-control/CLAUDE.md、CONTEXT.md、docs/OPEN_SOURCE.md 与本文件。")
numbered("认知：真实实现全部在 app/；packages/ 和 demo/ 是遗留，不以它们为入口。")
numbered("确认产品方向：独立 Android App，连接方式 STA / WiFi / USB，主路径 PTP/IP over WiFi。")
numbered("优先做真机验证：连接相机热点，启动 App，把「PTP/IP 握手诊断」日志发回。")
numbered("按日志校准 app/src/ptpip.js；协议未校准前不要假设厂商命令可用。")
numbered("改动后执行 npx vite build → npx cap sync android → cd android && ./gradlew assembleDebug。")
numbered("开源前执行 git status、git check-ignore，确认无 node_modules、APK、local.properties、密钥与参考 APK。")

heading("1. 交接目的与范围")
body("当前项目目标是做一款独立 Android 手机 App，控制和同步 Nikon Z30，并叠加 AI 修图与拍照姿势框线。真正的实现位于 app/，之前 packages/ 是早期蓝图，不要作为开发入口。")
bullet("主路径：PTP/IP over WiFi 无线遥控；USB 有线为备选。")
bullet("连接方式按参考 App 定义：STA、WiFi 热点、USB。")
bullet("AI 修图不锁死 DeepSeek，支持切换模型/接口。")
bullet("参考安装包只用于学习架构与交互，产物不得上传。")

heading("2. 当前项目状态")
table(
    ["模块", "状态", "说明"],
    [
        ["APK 构建", "已完成", "app-debug.apk 已能产出，包名 com.nikon.camera.control"],
        ["手机原生 TCP", "已实现", "自研 TcpSocketPlugin，可直连相机 15740"],
        ["PTP/IP 会话", "代码完成", "src/ptpip.js，握手需真机日志验证"],
        ["5 标签 UI", "已实现", "首页/我的相机/相机照片/同步/本地照片"],
        ["STA/WiFi/USB", "UI+路由完成", "WiFi/STA 走 TCP；USB 手机端待原生接入"],
        ["AI 修图", "已实现", "支持模型切换，图片可传视觉模型"],
        ["姿势框线", "已实现", "半身/全身/头部/三分/网格，透明度可调"],
    ],
    [1.25, 1.25, 4.0],
)

heading("3. 参考 App 已吸收内容")
body("参考安装包为 CameraSyncPro（Flutter）。已拆出的关键结构：")
bullet("底部导航：首页、我的相机、相机照片、同步、本地照片。")
bullet("连接详情：连接状态、上次连接方式（STA 局域网）、协议 PTP/IP、Wi-Fi 名。")
bullet("取景叠加层：live_view_overlays.dart、GridOverlay、GridPainter、liveViewGridLines。")
bullet("服务分层：camera_connection_service、wifi_camera_connection_service、ptp_ip_client、usb_remote_control_service、sync_task_manager。")

heading("4. 技术架构与关键目录")
table(
    ["路径", "作用"],
    [
        ["app/src/ptpip.js", "PTP/IP 协议与传输抽象"],
        ["app/src/api.js", "后端/原生双模式 API"],
        ["app/src/ai.js", "可切换模型 AI 客户端"],
        ["app/src/components/PoseGuideOverlay.jsx", "姿势框线叠加层"],
        ["app/src/components/BottomNav.jsx", "5 标签导航"],
        ["app/android/app/src/main/java/.../TcpSocketPlugin.java", "原生 TCP 插件"],
        ["app/server.cjs", "桌面/浏览器后端（USB、HTTP、WS）"],
    ],
    [3.0, 3.5],
)

heading("5. 连接方式")
table(
    ["模式", "场景", "实现现状"],
    [
        ["WiFi 热点", "手机连相机自带热点，默认 192.168.1.1", "已接入原生 TCP"],
        ["STA", "相机加入家庭 WiFi，手机同网，输入相机 IP", "已接入原生 TCP，IP 需手动输入"],
        ["USB", "手机 OTG 或电脑 USB 控制", "电脑后端可用；手机原生通道待接入"],
    ],
    [1.35, 2.75, 2.4],
)

heading("6. AI 修图")
bullet("默认 DeepSeek：https://api.deepseek.com/v1/chat/completions，模型 deepseek-chat。")
bullet("设置页可切换 DeepSeek / OpenAI 兼容 / 自定义模型、接口地址、模型名。")
bullet("支持传入图片 base64，视觉模型可收到图片；文字模型仍可给修图建议。")
bullet("API Key 仅存 localStorage，禁止写入仓库。")

heading("7. 人像拍照姿势框线")
bullet("实时取景与拍照时显示，可一键关闭。")
bullet("模式：半身、全身、头部特写、三分构图、对焦网格。")
bullet("透明度、线条颜色可调，设置持久化到 localStorage。")

heading("8. 构建与运行")
numbered("cd app && npm install")
numbered("npx vite build")
numbered("npx cap sync android")
numbered("cd android && ./gradlew assembleDebug")
numbered("输出：app/android/app/build/outputs/apk/debug/app-debug.apk")

heading("9. 本机工具链与注意")
bullet("JDK 21：C:\\Users\\李\\.jdks\\jdk-21.0.2（用 JAVA_HOME 传入，不要写进 gradle.properties）。")
bullet("Android SDK：C:\\Users\\Public\\AndroidSdk 是指向 C:\\Users\\李\\Android 的 ASCII junction。")
bullet("Windows 中文路径需 android.overridePathCheck=true。")
bullet("Gradle 8.14.3 zip 已缓存；以上均为本机信息，不上传。")

heading("10. 已知问题 / 待办")
table(
    ["优先级", "事项", "备注"],
    [
        ["P0", "真机验证 PTP/IP 握手", "把手机端握手诊断发回，校准 ptpip.js"],
        ["P1", "手机端 LiveView 数据阶段", "StartData/Data/EndData 尚未接入"],
        ["P1", "手机 USB 原生通道", "需要 USB Host 插件或 OTG 实现"],
        ["P2", "照片真实下载/本地相册", "当前本地媒体以文件选择为主"],
        ["P2", "Electron EXE 打包", "主进程已配置，未实际打包"],
        ["P3", "GitHub 开源", "已准备 .gitignore/README/LICENSE/docs/OPEN_SOURCE.md"],
    ],
    [0.8, 2.2, 3.5],
)

heading("11. GitHub 开源上传规则")
table(
    ["可以上传", "禁止上传"],
    [
        ["app/src/**、app/server.cjs、app/main-process/**", "node_modules、dist、build、.gradle"],
        ["app/android 源码与 Gradle Wrapper（除 jar 外）", "local.properties、*.apk、*.aab"],
        ["README.md、LICENSE、docs/", "packages、demo、CLAUDE/CONTEXT/AGENTS"],
        ["package.json、package-lock.json、配置文件", ".apk_ref、参考 APK、密钥、本机路径"],
    ],
    [3.25, 3.25],
)
body("详见 docs/OPEN_SOURCE.md。提交前执行 git status 与 git check-ignore 检查。", size=10, color="585870")

heading("12. 交接提示词（可整段复制给新对话）")
body("继续开发 Nikon Camera Control。先完整读 nikon-camera-control/CLAUDE.md、nikon-camera-control/CONTEXT.md、docs/OPEN_SOURCE.md，再读 docs/Nikon_Camera_Control_交接文档.docx。真实实现位于 app/；products 方向是独立 Android App，连接方式 STA/WiFi/USB，主路径 PTP/IP over WiFi，USB 备选；还要完成 AI 修图（模型可切换）与人像拍照姿势框线。第一优先级是用真机验证 PTP/IP 握手并校准 src/ptpip.js；第二优先级接入手机端 LiveView 数据阶段和 USB 原生通道。不要上传 node_modules、build、APK、local.properties、.npmrc 或参考 APK。", size=10.5)

doc.save(OUT)
print("saved:", OUT)
